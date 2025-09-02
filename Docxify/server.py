#!/usr/bin/env python3
from pathlib import Path
import io
import argparse
import markdown
from bs4 import BeautifulSoup
from fastapi import FastAPI, UploadFile, Form
from fastapi.responses import FileResponse, HTMLResponse
from fastapi.middleware.cors import CORSMiddleware
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import uvicorn

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        relationship_type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(u)
    c = OxmlElement("w:color")
    c.set(qn("w:val"), "0000FF")
    r_pr.append(c)
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def markdown_to_html(md_text: str) -> str:
    extensions = ["tables", "fenced_code", "footnotes", "attr_list", "sane_lists"]
    return markdown.markdown(md_text, extensions=extensions, output_format="html5")


def html_to_docx(html: str, doc: Document, base_path: Path):
    soup = BeautifulSoup(html, "html.parser")

    def handle_element(el, parent=None):
        if el.name is None:
            if parent:
                parent.add_run(el)
            return

        if el.name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
            level = int(el.name[1])
            doc.add_paragraph(el.get_text(), style=f"Heading {level}")

        elif el.name == "p":
            p = doc.add_paragraph()
            for child in el.children:
                handle_element(child, p)

        elif el.name in ["strong", "b"]:
            run = parent.add_run(el.get_text())
            run.bold = True

        elif el.name in ["em", "i"]:
            run = parent.add_run(el.get_text())
            run.italic = True

        elif el.name == "code" and el.parent.name != "pre":
            run = parent.add_run(el.get_text())
            run.font.name = "Courier New"

        elif el.name == "pre":
            p = doc.add_paragraph()
            run = p.add_run(el.get_text())
            run.font.name = "Courier New"
            run.font.size = Pt(10)

        elif el.name == "ul":
            for li in el.find_all("li", recursive=False):
                doc.add_paragraph(li.get_text(), style="List Bullet")

        elif el.name == "ol":
            for li in el.find_all("li", recursive=False):
                doc.add_paragraph(li.get_text(), style="List Number")

        elif el.name == "blockquote":
            doc.add_paragraph(el.get_text(), style="Intense Quote")

        elif el.name == "a":
            if parent:
                add_hyperlink(parent, el.get_text(), el.get("href", ""))

        elif el.name == "table":
            rows = el.find_all("tr")
            if rows:
                ncols = len(rows[0].find_all(["td", "th"]))
                table = doc.add_table(rows=len(rows), cols=ncols)
                for i, row in enumerate(rows):
                    cells = row.find_all(["td", "th"])
                    for j, cell in enumerate(cells):
                        table.cell(i, j).text = cell.get_text(strip=True)

        else:
            for child in el.children:
                handle_element(child, parent)

    for child in soup.children:
        handle_element(child)


def convert_markdown_to_docx(md_text: str, base_path: Path) -> io.BytesIO:
    html = markdown_to_html(md_text)
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    html_to_docx(html, doc, base_path)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


@app.get("/", response_class=HTMLResponse)
def index():
    return """
    <!DOCTYPE html>
    <html>
    <head>
        <title>Markdown to DOCX</title>
        <script src="https://cdn.tailwindcss.com"></script>
    </head>
    <body class="bg-gray-100 min-h-screen flex items-center justify-center">
      <div class="bg-white shadow-xl rounded-2xl p-8 w-full max-w-2xl">
        <h1 class="text-2xl font-bold mb-4 text-center text-gray-800">Markdown → DOCX</h1>
        <form id="uploadForm" class="space-y-4" enctype="multipart/form-data" method="post" action="/convert">
          <label class="block">
            <span class="text-gray-700">Upload a Markdown file</span>
            <input type="file" name="file" accept=".md" class="mt-2 block w-full text-sm text-gray-500 file:mr-4 file:py-2 file:px-4 file:rounded-xl file:border-0 file:text-sm file:font-semibold file:bg-blue-50 file:text-blue-700 hover:file:bg-blue-100"/>
          </label>
          <div class="text-center text-gray-500">or</div>
          <label class="block">
            <span class="text-gray-700">Paste Markdown text</span>
            <textarea name="text" rows="8" class="mt-2 block w-full border-gray-300 rounded-xl shadow-sm focus:ring-blue-500 focus:border-blue-500 sm:text-sm"></textarea>
          </label>
          <button type="submit" class="w-full py-2 px-4 bg-blue-600 text-white font-semibold rounded-xl shadow hover:bg-blue-700">Convert</button>
        </form>
      </div>
    </body>
    </html>
    """


@app.post("/convert")
async def convert(file: UploadFile = None, text: str = Form("")):
    if file:
        md_text = (await file.read()).decode("utf-8")
        base_path = Path(file.filename).parent
    elif text.strip():
        md_text = text
        base_path = Path(".")
    else:
        return {"error": "No input provided"}

    buffer = convert_markdown_to_docx(md_text, base_path)
    output_filename = "converted.docx"
    with open(output_filename, "wb") as f:
        f.write(buffer.getvalue())
    return FileResponse(
        output_filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=output_filename,
    )


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Markdown to DOCX FastAPI service")
    parser.add_argument("--host", type=str, default="0.0.0.0", help="Host address (default: 0.0.0.0)")
    parser.add_argument("--port", type=int, default=8000, help="Port (default: 8000)")
    args = parser.parse_args()

    uvicorn.run(app, host=args.host, port=args.port)
